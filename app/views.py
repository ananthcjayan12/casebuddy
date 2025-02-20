from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .forms import UserRegistrationForm
from .models import Case, Payment
from .forms import CaseForm, PaymentForm, CaseQuestionnaireForm
from django.urls import reverse
from decimal import Decimal
import openai  # You'll need to install this
from django.conf import settings
from django.http import HttpResponse
import docx
from docx.shared import Inches
import io

# Create your views here.

def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Registration successful. Please login.')
            return redirect('login')
    else:
        form = UserRegistrationForm()
    return render(request, 'registration/register.html', {'form': form})

@login_required
def profile(request):
    return render(request, 'app/profile.html')

@login_required
def submit_case(request):
    if request.user.userprofile.is_lawyer:
        messages.error(request, 'Lawyers cannot submit cases.')
        return redirect('case_list')
        
    if request.method == 'POST':
        form = CaseForm(request.POST)
        if form.is_valid():
            case = form.save(commit=False)
            case.client = request.user
            case.status = 'NEW'
            preferred_lawyer = form.cleaned_data.get('preferred_lawyer')
            if preferred_lawyer:
                case.lawyer = preferred_lawyer
            case.save()
            messages.success(request, 'Case submitted successfully. Please fill out the questionnaire.')
            return redirect('case_questionnaire', pk=case.id)
    else:
        form = CaseForm()
    
    return render(request, 'app/submit_case.html', {'form': form})

@login_required
def case_list(request):
    if request.user.userprofile.is_lawyer:
        cases = Case.objects.filter(lawyer=request.user)
    else:
        cases = Case.objects.filter(client=request.user)
    return render(request, 'app/case_list.html', {'cases': cases})

@login_required
def case_detail(request, pk):
    case = get_object_or_404(Case, pk=pk)
    
    # Check if user has permission to view this case
    if not (case.client == request.user or case.lawyer == request.user or 
            (request.user.userprofile.is_lawyer and not case.lawyer)):
        messages.error(request, 'You do not have permission to view this case.')
        return redirect('case_list')
    
    if request.method == 'POST' and request.user.userprofile.is_lawyer:
        new_status = request.POST.get('status')
        if new_status in dict(Case.STATUS_CHOICES):
            case.status = new_status
            case.save()
            messages.success(request, 'Case status updated successfully.')
            return redirect('case_detail', pk=pk)
    
    context = {
        'case': case,
        'status_choices': Case.STATUS_CHOICES,
        'payment_form': PaymentForm() if not request.user.userprofile.is_lawyer else None
    }
    return render(request, 'app/case_detail.html', context)

@login_required
def make_payment(request, case_id):
    case = get_object_or_404(Case, pk=case_id)
    
    # Verify that the user is the case client
    if case.client != request.user:
        messages.error(request, 'You are not authorized to make this payment.')
        return redirect('case_list')
    
    # Check if payment is already made
    if case.payment_status == 'COMPLETED':
        messages.warning(request, 'Payment has already been made for this case.')
        return redirect('case_detail', pk=case_id)
    
    if request.method == 'POST':
        form = PaymentForm(request.POST)
        if form.is_valid():
            payment = form.save(commit=False)
            payment.case = case
            payment.is_paid = True
            payment.save()
            
            # Update case status
            case.payment_status = 'COMPLETED'
            case.status = 'PENDING_PAYMENT'
            case.save()
            
            messages.success(request, 'Payment processed successfully!')
            return redirect('payment_success', payment_id=payment.id)
    else:
        # Pre-fill amount based on case type or default amount
        initial_amount = Decimal('500.00')  # You can modify this based on your needs
        form = PaymentForm(initial={'amount': initial_amount})
    
    return render(request, 'app/payment_form.html', {
        'case': case,
        'payment_form': form
    })

@login_required
def payment_success(request, payment_id):
    payment = get_object_or_404(Payment, pk=payment_id)
    
    # Verify that the user is the case client
    if payment.case.client != request.user:
        messages.error(request, 'You are not authorized to view this payment.')
        return redirect('case_list')
    
    return render(request, 'app/payment_success.html', {'payment': payment})

def home(request):
    return render(request, 'app/home.html')

@login_required
def available_cases(request):
    if not request.user.userprofile.is_lawyer:
        messages.error(request, 'Only lawyers can view available cases.')
        return redirect('case_list')
    
    # Get cases that don't have a lawyer assigned
    cases = Case.objects.filter(lawyer=None)
    return render(request, 'app/available_cases.html', {'cases': cases})

@login_required
def accept_case(request, pk):
    if not request.user.userprofile.is_lawyer:
        messages.error(request, 'Only lawyers can accept cases.')
        return redirect('case_list')
        
    case = get_object_or_404(Case, pk=pk)
    
    if case.lawyer:
        messages.error(request, 'This case has already been assigned to a lawyer.')
        return redirect('available_cases')
    
    case.lawyer = request.user
    case.status = 'ACTIVE'
    case.save()
    messages.success(request, f'You have successfully accepted the case: {case.title}')
    return redirect('case_detail', pk=pk)

@login_required
def case_questionnaire(request, pk):
    case = get_object_or_404(Case, pk=pk)
    
    if case.client != request.user:
        messages.error(request, 'You do not have permission to access this questionnaire.')
        return redirect('case_list')
    
    if hasattr(case, 'questionnaire'):
        messages.warning(request, 'Questionnaire has already been submitted.')
        return redirect('case_detail', pk=case.id)
    
    if request.method == 'POST':
        form = CaseQuestionnaireForm(request.POST)
        if form.is_valid():
            questionnaire = form.save(commit=False)
            questionnaire.case = case
            questionnaire.save()
            messages.success(request, 'Questionnaire submitted successfully.')
            return redirect('case_detail', pk=case.id)
    else:
        form = CaseQuestionnaireForm()
    
    return render(request, 'app/case_questionnaire.html', {'form': form, 'case': case})

@login_required
def generate_draft(request, pk):
    case = get_object_or_404(Case, pk=pk)
    
    # Only lawyers assigned to the case can generate drafts
    if not request.user.userprofile.is_lawyer or request.user != case.lawyer:
        messages.error(request, 'You do not have permission to generate drafts for this case.')
        return redirect('case_list')
    
    if not hasattr(case, 'questionnaire'):
        messages.error(request, 'Cannot generate draft - questionnaire not submitted yet.')
        return redirect('case_detail', pk=case.id)
    
    try:
        # Initialize OpenAI client with API key from settings
        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        
        # Prepare the prompt
        prompt = f"""
        Please prepare a legal case draft based on the following information:
        
        COURT INFORMATION:
        Court: {case.questionnaire.court_jurisdiction}
        Case Type: {case.get_category_display()}
        Case Number: {case.questionnaire.case_number if case.questionnaire.case_number else "To be assigned"}
        
        PLAINTIFF INFORMATION:
        Name: {case.questionnaire.client_name}
        Age: {case.questionnaire.client_age}
        Address: {case.questionnaire.client_address}
        Phone: {case.questionnaire.client_phone}
        Email: {case.questionnaire.client_email}
        
        DEFENDANT INFORMATION:
        Name: {case.questionnaire.defendant_name}
        Address: {case.questionnaire.defendant_address}
        Contact: {case.questionnaire.defendant_contact}
        
        INCIDENT DETAILS:
        Date: {case.questionnaire.incident_date}
        Location: {case.questionnaire.incident_location}
        
        Description:
        {case.questionnaire.incident_description}
        
        WITNESSES:
        {case.questionnaire.witnesses}
        
        PREVIOUS LEGAL ACTIONS:
        {case.questionnaire.previous_legal_actions}
        
        DESIRED OUTCOME:
        {case.questionnaire.desired_outcome}
        
        ADDITIONAL INFORMATION:
        {case.questionnaire.additional_info}
        
        Please format this into a professional legal document following standard court filing formats.
        Include all necessary sections such as jurisdiction statement, facts, legal grounds, and prayer for relief.
        """
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a legal professional drafting a case document."},
                {"role": "user", "content": prompt}
            ]
        )
        
        case.lawyer_draft = response.choices[0].message.content
        case.save()
        
        messages.success(request, 'Legal draft has been generated successfully.')
    except Exception as e:
        messages.error(request, f'Error generating draft: {str(e)}')
    
    return redirect('case_detail', pk=case.id)

@login_required
def download_draft(request, pk):
    case = get_object_or_404(Case, pk=pk)
    
    # Check permissions
    if not (case.client == request.user or case.lawyer == request.user):
        messages.error(request, 'You do not have permission to download this draft.')
        return redirect('case_list')
    
    if not case.lawyer_draft:
        messages.error(request, 'No draft available for download.')
        return redirect('case_detail', pk=pk)
    
    # Create a new Word document
    doc = docx.Document()
    
    # Add the content
    doc.add_heading('LEGAL DRAFT', 0)
    
    # Add case information
    doc.add_paragraph(f'Case Type: {case.get_category_display()}')
    doc.add_paragraph(f'Case Number: {case.questionnaire.case_number if case.questionnaire.case_number else "Not assigned"}')
    doc.add_paragraph(f'Court: {case.questionnaire.court_jurisdiction}')
    
    # Add a line
    doc.add_paragraph('_' * 50)
    
    # Add the main draft content
    doc.add_paragraph(case.lawyer_draft)
    
    # Save to a buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create the HTTP response
    response = HttpResponse(buffer.read())
    response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response['Content-Disposition'] = f'attachment; filename=legal_draft_{case.id}.docx'
    
    return response
